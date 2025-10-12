"""使用免费开源库将Markdown转换为Word、PDF和HTML"""

import logging
import os
from pathlib import Path

# 导入必要的库
try:
    import markdown
    MARKDOWN_AVAILABLE = True
    print("✅ markdown库导入成功")
except ImportError as e:
    MARKDOWN_AVAILABLE = False
    print(f"❌ markdown库导入失败: {e}")

try:
    from docx import Document  # type: ignore
    DOCX_AVAILABLE = True
    print("✅ python-docx库导入成功")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"❌ python-docx库导入失败: {e}")

try:
    import weasyprint  # type: ignore
    WEASYPRINT_AVAILABLE = True
    print("✅ weasyprint库导入成功")
    
    # 关闭WeasyPrint的详细日志
    logging.getLogger('weasyprint').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.document').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.css').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.html').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.text').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.layout').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.draw').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.fonts').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.progress').setLevel(logging.ERROR)
    logging.getLogger('weasyprint.utils').setLevel(logging.ERROR)
    
    # 关闭其他可能产生调试信息的日志
    logging.getLogger('fontTools').setLevel(logging.ERROR)
    logging.getLogger('fontTools.subset').setLevel(logging.ERROR)
    logging.getLogger('cffi').setLevel(logging.ERROR)
    
except ImportError as e:
    WEASYPRINT_AVAILABLE = False
    print(f"❌ weasyprint库导入失败: {e}")

try:
    import pdfkit  # type: ignore
    PDFKIT_AVAILABLE = True
    print("✅ pdfkit库导入成功")
except ImportError as e:
    PDFKIT_AVAILABLE = False
    print(f"❌ pdfkit库导入失败: {e}")


def print_installation_guide():
    """打印免费库安装指南"""
    print("\n=== 免费库安装指南 ===")
    print("以下都是免费开源库，可以替代Spire.Doc:")
    print()
    print("1. markdown (必需):")
    print("   pip install markdown")
    print()
    print("2. python-docx (Word转换):")
    print("   pip install python-docx")
    print()
    print("3. weasyprint (PDF转换):")
    print("   pip install weasyprint")
    print("   注意：Windows上可能需要额外配置")
    print()
    print("4. pdfkit (PDF转换备选):")
    print("   pip install pdfkit")
    print("   需要安装wkhtmltopdf: https://wkhtmltopdf.org/downloads.html")
    print()
    print("5. 如果都不安装，将使用HTML备选方案")
    print("=====================================\n")


def get_html_template(html_body: str) -> str:
    """获取HTML模板"""
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>顾问分析报告</title>
        <style>
            body {{
                font-family: "Microsoft YaHei", "SimSun", sans-serif;
                line-height: 1.6;
                margin: 2cm;
                color: #333;
            }}
            h1, h2, h3, h4, h5, h6 {{
                color: #2c3e50;
                margin-top: 1em;
                margin-bottom: 0.5em;
            }}
            h1 {{ font-size: 2.2em; border-bottom: 1px solid #eee; padding-bottom: 0.3em; }}
            h2 {{ font-size: 1.8em; }}
            h3 {{ font-size: 1.4em; }}
            h4 {{ font-size: 1.2em; }}
            p {{ margin-bottom: 1em; }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 1em;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #f2f2f2;
                font-weight: bold;
            }}
            pre {{
                background-color: #f6f8fa;
                padding: 1em;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 0.9em;
                margin-bottom: 1em;
            }}
            code {{
                font-family: 'Courier New', monospace;
                background-color: #f0f0f0;
                padding: 0.2em 0.4em;
                border-radius: 3px;
                font-size: 0.9em;
            }}
            blockquote {{
                border-left: 4px solid #ccc;
                padding-left: 1em;
                margin-left: 0;
                color: #555;
                font-style: italic;
                margin-bottom: 1em;
            }}
        </style>
    </head>
    <body>
        {html_body}
    </body>
    </html>
    """


def markdown_to_word_free(markdown_file: str, output_file: str) -> bool:
    """使用python-docx将Markdown转换为Word文档"""
    if not DOCX_AVAILABLE or not MARKDOWN_AVAILABLE:
        print("❌ 缺少必要的库 (python-docx 或 markdown)")
        return False
    
    try:
        # 读取Markdown文件
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # 转换为HTML（用于参考，但这里我们直接处理Markdown文本）
        # html_body = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        
        # 创建Word文档
        doc = Document()
        
        # 简单的文本处理 - 将HTML标签转换为Word格式
        lines = markdown_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)
        
        doc.save(output_file)
        print(f"✅ Markdown转Word成功: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ Markdown转Word失败: {e}")
        return False


def markdown_to_pdf_weasyprint(markdown_file: str, output_file: str) -> bool:
    """使用WeasyPrint将Markdown转换为PDF"""
    if not WEASYPRINT_AVAILABLE or not MARKDOWN_AVAILABLE:
        print("❌ 缺少必要的库 (weasyprint 或 markdown)")
        return False
    
    try:
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        html_body = markdown.markdown(markdown_content, extensions=['extra', 'codehilite', 'tables'])
        full_html = get_html_template(html_body)
        
        # 使用WeasyPrint转换，添加更多配置选项
        html_doc = weasyprint.HTML(string=full_html)
        css = weasyprint.CSS(string='''
            @page {
                size: A4;
                margin: 2cm;
            }
            body {
                font-family: "Microsoft YaHei", "SimSun", sans-serif;
                line-height: 1.6;
                color: #333;
            }
        ''')
        
        html_doc.write_pdf(output_file, stylesheets=[css])
        print(f"✅ Markdown转PDF成功 (WeasyPrint): {output_file}")
        return True
        
    except OSError as e:
        if "libgobject" in str(e) or "cannot load library" in str(e):
            print(f"❌ WeasyPrint遇到GTK依赖问题: {e}")
            print("💡 请确保已正确安装GTK3和相关依赖")
            return False
        else:
            print(f"❌ WeasyPrint系统依赖错误: {e}")
            return False
    except Exception as e:
        print(f"❌ WeasyPrint转换失败: {e}")
        return False


def markdown_to_pdf_pdfkit(markdown_file: str, output_file: str) -> bool:
    """使用pdfkit将Markdown转换为PDF"""
    if not PDFKIT_AVAILABLE or not MARKDOWN_AVAILABLE:
        print("❌ 缺少必要的库 (pdfkit 或 markdown)")
        return False
    
    try:
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        html_body = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        full_html = get_html_template(html_body)
        
        options = {
            'page-size': 'A4',
            'margin-top': '2cm',
            'margin-right': '2cm',
            'margin-bottom': '2cm',
            'margin-left': '2cm',
            'encoding': "UTF-8",
            'no-outline': None
        }
        
        pdfkit.from_string(full_html, output_file, options=options)
        print(f"✅ Markdown转PDF成功 (pdfkit): {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ pdfkit转换失败: {e}")
        return False


def markdown_to_html_free(markdown_file: str, output_file: str) -> bool:
    """将Markdown转换为HTML文档"""
    if not MARKDOWN_AVAILABLE:
        print("❌ 缺少markdown库")
        return False
    
    try:
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        html_body = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        full_html = get_html_template(html_body)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        print(f"✅ Markdown转HTML成功: {output_file}")
        return True
        
    except Exception as e:
        print(f"❌ Markdown转HTML失败: {e}")
        return False


def test_free_conversion():
    """测试免费库转换功能"""
    
    if not MARKDOWN_AVAILABLE:
        print_installation_guide()
        return
    
    # 检查可用的转换库
    available_conversions = []
    if DOCX_AVAILABLE:
        available_conversions.append("Word")
    if WEASYPRINT_AVAILABLE or PDFKIT_AVAILABLE:
        available_conversions.append("PDF")
    available_conversions.append("HTML")
    
    print(f"✅ 可用的转换格式: {', '.join(available_conversions)}")
    
    # 测试文件列表
    test_files = [
        "consultant_analysis_output/2025-10-10/智浩_顾问分析报告_2025-10-10_20251010_184425.md",
        "consultant_analysis_output/2025-10-10/思怡_顾问分析报告_2025-10-10_20251010_184529.md", 
        "consultant_analysis_output/2025-10-10/佳慧_顾问分析报告_2025-10-10_20251010_184633.md"
    ]
    
    print("=== 开始测试免费库转换功能 ===")
    
    # 创建输出目录
    output_dir = Path("free_output")
    output_dir.mkdir(exist_ok=True)
    
    for md_file in test_files:
        if not os.path.exists(md_file):
            print(f"❌ 文件不存在: {md_file}")
            continue
            
        print(f"\n📄 正在转换: {md_file}")
        
        # 获取文件名（不含扩展名）
        file_name = Path(md_file).stem
        
        # 转换到Word
        if DOCX_AVAILABLE:
            word_file = output_dir / f"{file_name}.docx"
            markdown_to_word_free(md_file, str(word_file))
        
        # 转换到PDF
        if WEASYPRINT_AVAILABLE:
            pdf_file = output_dir / f"{file_name}_weasyprint.pdf"
            if not markdown_to_pdf_weasyprint(md_file, str(pdf_file)) and PDFKIT_AVAILABLE:
                pdf_file = output_dir / f"{file_name}_pdfkit.pdf"
                markdown_to_pdf_pdfkit(md_file, str(pdf_file))
        elif PDFKIT_AVAILABLE:
            pdf_file = output_dir / f"{file_name}.pdf"
            markdown_to_pdf_pdfkit(md_file, str(pdf_file))
        
        # 转换到HTML
        html_file = output_dir / f"{file_name}.html"
        markdown_to_html_free(md_file, str(html_file))
    
    print("\n=== 免费库转换测试完成 ===")


if __name__ == "__main__":
    test_free_conversion()
